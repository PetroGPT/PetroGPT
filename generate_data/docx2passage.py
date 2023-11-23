'''
将一堆文档切分成段落。
输入：docx目录
输出：json文件
    字段： 段落id，段落passage，长度len，路径fpath，起始start end
'''
import argparse
import os.path
import re
import glob
from docx import Document
from dalchemy.data import DalchemyData,TextHelper
from dalchemy.llms import AzureLLM,OpenAILLM
from langchain.text_splitter import RecursiveCharacterTextSplitter,SpacyTextSplitter
# from pipelines.nodes import CharacterTextSplitter
helper = TextHelper()
from tqdm import tqdm
from docx import Document as DP
import zhconv

def has_non_alphanumeric(text, threshold=0.5):
    # 移除空格和标点符号
    cleaned_text = re.sub(r'\s+|[^\w\s]', '', text)
    if cleaned_text.strip() == "": return False

    # 计算数字和英文字母的比例
    total_chars = len(cleaned_text)
    total_digits = len(re.findall(r'\d', cleaned_text))
    total_letters = len(re.findall(r'[a-zA-Z]', cleaned_text))
    ratio = (total_digits + total_letters) / total_chars

    # 判断比例是否超过阈值
    print("ratio",ratio)
    if ratio > threshold:
        return False
    else:
        return True

def has_alphanumeric(text, threshold=0.5):
    # 移除空格和标点符号
    cleaned_text = re.sub(r'\s+|[^\w\s]', '', text)
    if cleaned_text.strip() == "": return False

    # 计算数字和英文字母的比例
    total_chars = len(cleaned_text)
    total_digits = len(re.findall(r'\d', cleaned_text))
    total_letters = len(re.findall(r'[a-zA-Z]', cleaned_text))
    ratio = (total_digits + total_letters) / total_chars

    # 判断比例是否超过阈值
    if ratio > threshold:
        return True
    return False

def read_pydocx(file_path):
    res = []
    docx_file = DP(file_path)
    for paragraph in docx_file.paragraphs:
        if not paragraph.text.strip():
            continue
        text = zhconv.convert(paragraph.text, 'zh-cn')
        res.append(text)
    return res

def get_doc_name(file_path):
    doc_name = os.path.basename(file_path)

    doc_name_regex = re.compile(r'^(?P<doc>.*?)_[\d]+【.*.docx$')
    m_d = doc_name_regex.match(doc_name)
    if m_d:
        doc_name = m_d["doc"]
    doc_name = doc_name.replace(".docx", "")
    return doc_name

def convert(rule, chunk_size=600, overlap=50, ret_pt_data = False):
    # text_splitter = CharacterTextSplitter(separator="\n", chunk_size=chunk_size, chunk_overlap=0, filters=["\n"])
    text_splitter = RecursiveCharacterTextSplitter(
        # Set a really small chunk size, just to show.
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        length_function=len,
        add_start_index=True,
    )

    files = glob.glob(rule)
    res = []
    res_pt = []
    idx = 0
    seen = set()
    for file in tqdm(files):
        filename = os.path.basename(file)
        if filename.startswith("~$"): continue
        doc = Document(file)
        text_ls = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text == "": continue
            text_ls.append(text)
        paragraph = "\n".join(text_ls)
        text_ls = text_splitter.create_documents([paragraph])
        for text in text_ls:
            content = text.page_content.strip()
            if has_alphanumeric(content):
                print("字符太多：", content,"--------------------------------------\n")
                continue

            if content not in seen:
                seen.add(content)
            else:
                continue
            # start =
            record = {"pid": str(idx),
                      "len": len(content),
                      "fpath":file,
                      "content":content}
            pt_data = {"text": content}
            res.append(record)
            res_pt.append(pt_data)
            idx += 1
    if ret_pt_data:
        helper.write_jsonl(res_pt, "text_data.jsonl")
    return res


def main():
    parser = argparse.ArgumentParser(description="将一堆文档切分成段落。")
    parser.add_argument('--rule', required=True, type=str, help='要搜索的文档的路径规则，如 "folder/*.docx"')
    parser.add_argument('--chunk_size',required=True, type=int, default=1024, help='处理文档时使用的块大小，默认为1024')
    parser.add_argument('--overlap', required=True, type=int, default=50, help='块之间的重叠大小，默认为50')
    parser.add_argument('--ret_pt_data', required=True, action='store_true', help='是否返回处理后的数据，默认不返回')
    parser.add_argument('--outfile', required=True,type=str, default="data/passages_all.json", help='输出JSON文件的路径，默认为 "data/passages_all.json"')

    args = parser.parse_args()

    res = convert(args.rule, args.chunk_size, args.overlap, args.ret_pt_data)
    print("段落数量：", len(res))
    helper.write_json(res, json_outfile=args.outfile)


if __name__ == '__main__':
    main()